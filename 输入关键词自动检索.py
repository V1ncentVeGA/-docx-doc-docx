# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches
import sys
import pickle
import re
import codecs
import string
import shutil
from win32com import client as wc
import docx


def doSaveAas(path,title):
    word = wc.Dispatch('Word.Application')
    mid_path = path + '\\' + title
    # print(mid_path)
    final_path = path + '\\' + title + 'x'
    # print(final_path)
    doc = word.Documents.Open(mid_path)  # 目标路径下的文件
    doc.SaveAs(final_path, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()



current_directory = os.path.dirname(os.path.abspath(__file__))
path = current_directory #文件夹目录
files= os.listdir(path) #得到文件夹下的所有文件名称/
# print(path)
# print(files)

# doSaveAas(path,files[0])

s = []
titles = []
for file in files: #遍历文件夹
    if not os.path.isdir(file): #判断是否是文件夹，不是文件夹才打开
        if ".docx" in file:
            document = Document(file)  #打开文件 *.docx
            str = ""
            for paragraph in document.paragraphs:
                str = str + paragraph.text  # 打印各段落内容文本
            document.save(file)  # 保存文档
            titles.append(file)
            s.append(str)  # 每个文件的文本存到list中

final_list = list(zip(titles,s))
# print(final_list) #打印结果

flag = True
keyword = input("请输入关键词：")
print("*"*18)
print()
for i in final_list :
    if keyword in i[1] :
        print(i[0])
        flag = False
print()
print("*"*18)
if flag :
    print("没有找到此关键词")

input("输入任意键两次退出")
input()
        # with open(path+"/"+file, encoding='UTF-8') as f :#打开文件
        #     iter_f = iter(f)#创建迭代器
        #     str = ""
        #     for line in iter_f: #遍历文件，一行行遍历，读取文本
        #         str = str + line
        #         s.append(str) #每个文件的文本存到list中


# document = Document('demo.docx')  #打开文件demo.docx
# for paragraph in document.paragraphs:
#     print(paragraph.text)  #打印各段落内容文本
# document.add_paragraph(
#     'Add new paragraph', style='ListNumber'
# )    #添加新段落
# document.save('demo.docx') #保存文档